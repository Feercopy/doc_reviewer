from pathlib import Path

from docx import Document


def parse_docx(path: Path) -> str:
    document = Document(path)
    blocks: list[str] = []

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            blocks.append(text)

    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                blocks.append("\t".join(cells))

    return "\n\n".join(blocks)
